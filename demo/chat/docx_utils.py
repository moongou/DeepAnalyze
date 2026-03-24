"""
DOCX 中文生成工具模块

此模块专门解决 DOCX 文件生成时的中文显示问题：
1. 正确设置中文字体
2. 确保使用 UTF-8 编码
3. 避免乱码问题
"""

import os
import re
import traceback
import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ==================== 字体配置 ====================
# 使用中央配置模块
try:
    from config import get_fonts_dir
    FONTS_DIR = get_fonts_dir()
except ImportError:
    # 回退到相对路径解析
    _CURRENT_DIR = Path(__file__).parent.resolve()
    _PROJECT_ROOT = _CURRENT_DIR.parent
    FONTS_DIR = _PROJECT_ROOT / "assets" / "fonts"

# 验证字体目录是否存在
if not FONTS_DIR.exists():
    logger.error(f"字体目录不存在: {FONTS_DIR}")
    # 尝试备用路径
    _ALTERNATIVE_FONTS = [
        Path(__file__).parent.parent / "assets" / "fonts",
        Path.home() / "DeepAnalyze" / "assets" / "fonts",
    ]
    for alt_path in _ALTERNATIVE_FONTS:
        if alt_path.exists():
            FONTS_DIR = alt_path
            logger.info(f"使用备用字体目录: {FONTS_DIR}")
            break
    if not FONTS_DIR.exists():
        logger.error("请检查字体文件是否存在于 assets/fonts/ 目录")


def get_font_dir() -> Path:
    """获取字体目录"""
    return FONTS_DIR


def get_chinese_font_path() -> Optional[str]:
    """
    获取可用的中文字体路径

    返回:
        str: 字体文件路径，如果找不到则返回 None
    """
    # 优先使用 assets/fonts 目录下的纯 TTF 字体
    font_files = [
        ("STFangSong.ttf", "仿宋 - 正文/报告"),
        ("SimHei.ttf", "黑体 - 主标题"),
        ("SimKai.ttf", "楷体 - 强调"),
        ("STHeiti.ttf", "黑体备选"),
        ("LiSongPro.ttf", "隶书"),
    ]

    for font_name, description in font_files:
        font_path = FONTS_DIR / font_name
        if font_path.exists():
            logger.info(f"找到中文字体: {font_path} ({description})")
            return str(font_path.resolve())

    logger.warning(f"未找到中文字体，字体目录: {FONTS_DIR}")
    return None


def create_docx_document() -> tuple[Document, Optional[str]]:
    """
    创建配置好中文字体的 DOCX 文档

    Returns:
        tuple: (doc, font_name) - 文档对象和字体名称
    """
    doc = Document()

    # 获取中文字体路径
    font_path = get_chinese_font_path()

    # 设置默认样式
    style = doc.styles["Normal"]
    if font_path:
        # 使用 STFangsong（注意：实际字体名称是小写 g）
        # DOCX 格式区分大小写，必须使用正确的字体名称
        font_name = "STFangsong"  # 仿宋 - 正确的注册名称
        style.font.name = font_name
        rFonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(11)
    else:
        font_name = None
        style.font.name = "Arial"

    # 设置标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f"Heading {i}"]
        if font_name:
            heading_style.font.name = font_name
            rFonts = heading_style._element.get_or_add_rPr().get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), font_name)
            heading_style.font.bold = True

            # 设置字号
            if i == 1:
                heading_style.font.size = Pt(16)
            elif i == 2:
                heading_style.font.size = Pt(14)
            else:
                heading_style.font.size = Pt(12)

    return doc, font_name


def clean_md_text_for_docx(md_text: str) -> str:
    """
    清理 Markdown 文本，移除 DOCX 不支持的标记

    Args:
        md_text: 原始 Markdown 文本

    Returns:
        str: 清理后的文本
    """
    if not md_text:
        return ""

    # 移除 \newpage 标记
    text = re.sub(r'\\newpage', '', md_text)
    # 移除 LaTeX 标记
    text = re.sub(r'\\[a-zA-Z]+\{[^}]*\}', '', text)
    # 移除 HTML 注释
    text = re.sub(r'<!--[^>]*-->', '', text)

    return text.strip()


def extract_markdown_blocks(md_text: str) -> list[dict]:
    """
    从 Markdown 文本中提取块结构

    Args:
        md_text: Markdown 格式的文本

    Returns:
        list: 块列表，每个块包含 type 和 content
    """
    if not md_text:
        return []

    blocks = []
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 图像检测
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
        if img_match:
            blocks.append({
                "type": "image",
                "content": img_match.group(2),
                "alt": img_match.group(1),
            })
            i += 1
            continue

        # 一级标题
        if stripped.startswith('# '):
            blocks.append({
                "type": "heading1",
                "content": stripped[2:].strip(),
            })
            i += 1
            continue

        # 二级标题
        if stripped.startswith('## '):
            blocks.append({
                "type": "heading2",
                "content": stripped[3:].strip(),
            })
            i += 1
            continue

        # 三级标题
        if stripped.startswith('### '):
            blocks.append({
                "type": "heading3",
                "content": stripped[4:].strip(),
            })
            i += 1
            continue

        # 列表项
        if stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.\s', stripped):
            list_items = []
            while i < len(lines):
                line = lines[i].strip()
                if line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.\s', line):
                    # 移除列表标记
                    content = re.sub(r'^[-*]\s+', '', line)
                    content = re.sub(r'^\d+\.\s+', '', content)
                    list_items.append(content)
                    i += 1
                elif not line:
                    i += 1
                    continue
                else:
                    break

            if list_items:
                blocks.append({
                    "type": "list",
                    "content": list_items,
                })
            continue

        # 普通段落
        paragraph_lines = []
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if not stripped or stripped.startswith(('# ', '## ', '### ', '- ', '* ')) or re.match(r'^\d+\.\s', stripped):
                break
            paragraph_lines.append(line)
            i += 1

        if paragraph_lines:
            blocks.append({
                "type": "paragraph",
                "content": '\n'.join(paragraph_lines),
            })

    return blocks


def add_block_to_docx(doc: Document, block: dict, font_name: Optional[str] = None) -> None:
    """
    将块添加到 DOCX 文档

    Args:
        doc: DOCX 文档对象
        block: 块字典
        font_name: 字体名称
    """
    block_type = block["type"]
    content = block["content"]

    if block_type == "heading1":
        p = doc.add_heading(level=1)
        run = p.add_run(content)
        if font_name:
            run.font.name = font_name
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    elif block_type == "heading2":
        p = doc.add_heading(level=2)
        run = p.add_run(content)
        if font_name:
            run.font.name = font_name
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    elif block_type == "heading3":
        p = doc.add_heading(level=3)
        run = p.add_run(content)
        if font_name:
            run.font.name = font_name
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    elif block_type == "image":
        try:
            img_path = content
            # 同样尝试解析路径
            if not os.path.isabs(img_path):
                # 如果是相对路径，尝试在 doc 的保存目录寻找 (如果尚未保存，output_path 会由 generate_docx 传入)
                # 这种情况下，我们假设 images 和 doc 在同一目录或相对
                pass # 路径由调用者处理或由 os.path 决定

            if os.path.exists(img_path):
                from docx.shared import Inches
                doc.add_picture(img_path, width=Inches(6.0)) # 默认宽度 6 英寸
                if block.get("alt"):
                    p = doc.add_paragraph(block["alt"])
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if font_name:
                        for run in p.runs:
                            run.font.name = font_name
                            run.font.italic = True
                            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)
            else:
                logger.warning(f"DOCX 图像缺失: {img_path}")
                p = doc.add_paragraph(f"[图像缺失: {img_path}]")
        except Exception as e:
            logger.error(f"DOCX 添加图像失败: {e}")

    elif block_type == "paragraph":
        p = doc.add_paragraph(content)
        if font_name:
            for run in p.runs:
                run.font.name = font_name
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    elif block_type == "list":
        for item in content:
            p = doc.add_paragraph(f"• {item}", style="List Paragraph")
            if font_name:
                for run in p.runs:
                    run.font.name = font_name
                    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)


def generate_docx(
    md_text: str,
    output_path: str,
    title: Optional[str] = None
) -> bool:
    """
    从 Markdown 文本生成 DOCX 文件

    Args:
        md_text: Markdown 格式的文本内容
        output_path: 输出 DOCX 文件路径
        title: 文档标题（可选）

    Returns:
        bool: 生成成功返回 True，否则返回 False
    """
    try:
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        logger.info(f"开始生成 DOCX: {output_path}")

        # 1. 创建文档并设置字体
        doc, font_name = create_docx_document()
        if font_name:
            logger.info(f"使用字体: {font_name}")
        else:
            logger.warning("未找到中文字体，DOCX 可能显示乱码")

        # 2. 清理 Markdown 文本
        clean_text = clean_md_text_for_docx(md_text)

        # 3. 提取内容块
        blocks = extract_markdown_blocks(clean_text)
        logger.info(f"提取到 {len(blocks)} 个内容块")

        # 4. 添加内容到文档
        for block in blocks:
            if block["type"] == "image":
                img_path = block["content"]
                if not os.path.isabs(img_path):
                    full_img_path = os.path.join(output_dir, img_path)
                    if not os.path.exists(full_img_path):
                        # 如果还没找到，尝试当前目录
                        full_img_path = os.path.abspath(img_path)
                    block["content"] = full_img_path
            add_block_to_docx(doc, block, font_name)

        # 5. 保存文档
        doc.save(output_path)

        # 6. 验证生成结果
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            file_size = os.path.getsize(output_path) / 1024
            logger.info(f"DOCX 生成成功: {output_path} ({file_size:.1f} KB)")
            return True
        else:
            logger.error(f"DOCX 生成失败: 文件不存在或为空")
            return False

    except Exception as e:
        error_msg = f"DOCX 生成失败: {e}"
        logger.error(error_msg)
        traceback.print_exc()
        return False


def test_docx_generation() -> str:
    """
    测试 DOCX 生成功能

    Returns:
        str: 测试结果路径或错误信息
    """
    import tempfile

    test_md = """# 测试报告

这是一个中文 DOCX 生成测试报告。

## 第一部分：字体测试

以下内容应该使用中文字体显示：

- 仿宋（STFangSong）用于正文
- 黑体（SimHei）用于标题
- 楷体（SimKai）用于强调

### 子测试：样式验证

验证不同级别的标题是否正确显示：

1. 一级标题应该较大且加粗
2. 二级标题应该适中且加粗
3. 三级标题应该较小且加粗

## 第二部分：段落测试

这是一段正常的正文内容，应该使用仿宋字体显示。

这是第二段正文，用于测试段落间距和行距是否正确。

## 第三部分：列表测试

以下是项目列表：

- 项目 1
- 项目 2
- 项目 3

以及编号列表：

1. 第一项
2. 第二项
3. 第三项

# 测试完成

如果以上内容都能正确显示中文，则说明字体设置和 DOCX 生成功能正常。
"""

    # 创建临时测试文件
    temp_dir = tempfile.gettempdir()
    test_output = os.path.join(temp_dir, "test_chinese_docx.docx")

    success = generate_docx(test_md, test_output, title="中文 DOCX 测试报告")

    if success:
        return test_output
    else:
        return "测试失败，请查看日志了解详情"


if __name__ == "__main__":
    print("=" * 60)
    print("开始测试中文 DOCX 生成...")
    print("=" * 60)

    result = test_docx_generation()

    if os.path.exists(result):
        print(f"\n✓ 测试成功！DOCX 已生成: {result}")
        print(f"  文件大小: {os.path.getsize(result) / 1024:.1f} KB")
    else:
        print(f"\n✗ 测试失败: {result}")

    print("=" * 60)
